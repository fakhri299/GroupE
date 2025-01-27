from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.platypus import Spacer
from reportlab.platypus.flowables import PageBreak
from reportlab.lib.styles import getSampleStyleSheet
import os
import docx


download_dir=os.path.expanduser("~/Downloads")

#FM-create_pdf creates report as pdf file

def create_pdf(db_links,cve_id, summary, cvss_score, vector, reference1,reference2):
    pdf_filename = f"{download_dir}/{cve_id}.pdf"

    doc = SimpleDocTemplate(pdf_filename, pagesize=letter)
    story = []

    # Define a style for the paragraphs
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]

    # Add information to the PDF with line breaks after each header
    story.append(Paragraph(f"<b>CVE ID:\n</b>", normal_style))
    story.append(Paragraph(f"{cve_id}",normal_style))
    story.append(Spacer(1, 12))
    if len(db_links)>=1:
        story.append(Paragraph(f"<b>Exploit found on Exploit-DB:</b>", normal_style))
        for item in db_links:
            story.append(Paragraph(item, normal_style))
        story.append(Spacer(1, 12))
    else:
        story.append(Paragraph(f"<b>Exploit not found on Exploit-DB</b>", normal_style))
        story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Summary:\n</b>", normal_style))
    story.append(Paragraph(f"{summary}",normal_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Vector:\n</b>", normal_style))
    story.append(Paragraph(f"{vector}", normal_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Base Score:\n</b>", normal_style))
    story.append(Paragraph(f"{cvss_score}",normal_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>References:\n</b>", normal_style))
    story.append(Paragraph(f"{reference1}",normal_style))
    story.append(Paragraph(f"{reference2}",normal_style))
    story.append(Spacer(1, 12))

    # Add a page break to start a new page for each CVE report
    story.append(PageBreak())

    # Save the PDF
    doc.build(story)

#FM-create_txt creates report as txt file
def create_txt(cve_id, summary, cvss_score, vector, reference1, reference2, db_links):
    txt_filename = f"{download_dir}/{cve_id}.txt"  # Specify the desired file path

    with open(txt_filename, 'w') as txt_file:
        # Write the CVE ID to the text file
        txt_file.write(f"CVE ID:\n{cve_id}\n\n")

        if db_links and len(db_links) >= 1:
            # Write Exploit-DB links if available
            txt_file.write("Exploit found on Exploit-DB:\n")
            for item in db_links:
                txt_file.write(f"{item}\n")
            txt_file.write("\n")
        else:
            # Write a message if no Exploit-DB links are available
            txt_file.write("Exploit not found on Exploit-DB\n\n")

        # Write the Summary to the text file
        txt_file.write(f"Summary:\n{summary}\n\n")

        # Write the Vector to the text file
        txt_file.write(f"Vector:\n{vector}\n\n")

        # Write the Base Score to the text file
        txt_file.write(f"Base Score:\n{cvss_score}\n\n")

        # Write the References to the text file
        txt_file.write("References:\n")
        txt_file.write(f"{reference1}\n")
        txt_file.write(f"{reference2}\n")

#SA-create_docx creates report as docx file

def create_docx(cve_id, summary, cvss_score, vector, reference1, reference2, db_links=None):
    docx_filename = f"{download_dir}/{cve_id}.docx"  # Specify the desired file path

    # Create a new DOCX document
    doc = docx.Document()

    # Add a title
    doc.add_heading('CVE Report', 0)

    # Add CVE ID
    doc.add_heading('CVE ID', level=1)
    doc.add_paragraph(cve_id)

    if db_links and len(db_links) >= 1:
        # Add Exploit-DB links if available
        doc.add_heading('Exploit found on Exploit-DB', level=1)
        for item in db_links:
            doc.add_paragraph(item)
    else:
        # Add a message if no Exploit-DB links are available
        doc.add_heading('Exploit not found on Exploit-DB', level=1)

    # Add Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)

    # Add Vector
    doc.add_heading('Vector', level=1)
    doc.add_paragraph(vector)

    # Add Base Score
    doc.add_heading('Base Score', level=1)
    doc.add_paragraph(cvss_score)

    # Add References
    doc.add_heading('References', level=1)
    doc.add_paragraph(reference1)
    doc.add_paragraph(reference2)

    # Save the DOCX document
    doc.save(docx_filename)


import os
import zipfile

#NR-compress() compresses report created and script downloaded to zip file
def compress(directory):
    # Get a list of all files in the directory
    all_files = os.listdir(directory)

    # Sort the files by modification time to get the top 2
    all_files.sort(key=lambda x: os.path.getmtime(os.path.join(directory, x)), reverse=True)
    top_two_files = all_files[:2]

    if len(top_two_files) < 2:
        print("There are not enough files to create a zip archive.")
        return

    # Create a zip file with a generic name
    zip_filename = os.path.join(directory, 'top_two_files.zip')

    with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file in top_two_files:
            file_path = os.path.join(directory, file)
            zipf.write(file_path, os.path.basename(file_path))

